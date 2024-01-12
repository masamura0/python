import os
import sys
import pandas as pd 
from docx import Document
path = os.getcwd()
try:
    df = pd.read_excel("list-for_tentant-onboarding.xlsx")
except FileNotFoundError:
    print("File not found")  
    sys.exit()
columns = list(df.columns) 
for index,item in df.iterrows(): 
    try:
        document = Document("template_tenant-onboarding_MOP.docx")
    except FileNotFoundError: 
        print("File not found") 
        sys.exit()
    for paragraph in document.paragraphs: 
        for column in columns: 
            if str(column) in paragraph.text: 
                def replace_text(paragraph):
                    replaced_text = paragraph.text.replace(str(column),str(item[column]))
                    if paragraph.text != replaced_text:
                        paragraph.text = replaced_text
                for paragraph in document.paragraphs:
                    replace_text(paragraph)
                paragraphs = (paragraph
                              for table in document.tables
                              for row in table.rows
                              for cell in row.cells
                              for paragraph in cell.paragraphs)
                for paragraph in paragraphs:
                    replace_text(paragraph)                   
    outfile = os.path.join(path,"old_tenant-onboarding",item["$file_name"]+".docx")
    try:
        document.save(outfile) 
    except FileNotFoundError:
        print("No destination folder")
        sys.exit()
